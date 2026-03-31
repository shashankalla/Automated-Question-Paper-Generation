import streamlit as st
from langchain_groq import ChatGroq
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
import os
import docx
from io import BytesIO
from dotenv import load_dotenv
import pdfplumber

# -------------------- ENV SETUP --------------------
load_dotenv()
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")

st.set_page_config(
    page_title="CO Based Question Generator",
    layout="wide"
)

if not GROQ_API_KEY or GROQ_API_KEY.strip() == "":
    st.error("❌ GROQ_API_KEY not found.")
    st.stop()

llm = ChatGroq(
    model_name="llama-3.1-8b-instant",
    temperature=0.5
)

# -------------------- PDF READER --------------------
def extract_text_from_pdf(uploaded_file):
    text = ""
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()

# -------------------- QUESTION GENERATOR --------------------
def generate_questions(
    q_type,
    subject,
    full_syllabus,
    unit_number,
    co_name,
    number,
    example
):
    if number == 0:
        return ""

    prompt_template = """
    You are an expert university question paper setter.

    Generate {number} {q_type} questions.

    Subject: {subject}
    Course Outcome: {co_name}

    FULL SYLLABUS (from PDF):
    {full_syllabus}

    IMPORTANT:
    - Generate questions ONLY from Unit {unit_number}
    - Ignore all other units completely

    Example format:
    {example}

    Rules:
    - University exam standard
    - NO answers
    - NO explanations

    Output format:
    Q1. ...
    Q2. ...
    """

    chain = (
        ChatPromptTemplate.from_template(prompt_template)
        | llm
        | StrOutputParser()
    )

    return chain.invoke({
        "number": number,
        "q_type": q_type,
        "subject": subject,
        "co_name": co_name,
        "full_syllabus": full_syllabus,
        "unit_number": unit_number,
        "example": example
    })

# -------------------- SESSION STATE --------------------
for key in ["mcq", "short", "long", "selected_co", "unit", "syllabus_text"]:
    if key not in st.session_state:
        st.session_state[key] = ""

# -------------------- DARK NEON UI CSS --------------------
st.markdown("""
    <style>

    body, .stApp {
        background: radial-gradient(circle at top, #0a0f24, #000000 70%);
        color: #e3e3e3;
        font-family: 'Segoe UI', sans-serif;
    }

    .title-glow {
        font-size: 42px !important;
        font-weight: 700;
        text-align: center;
        color: #7f5cff;
        text-shadow: 0 0 20px #7f5cff;
        margin-bottom: 10px;
    }

    .subtitle {
        text-align: center;
        font-size: 18px;
        color: #9b8cff;
        margin-bottom: 40px;
    }

    .glass-card {
        background: rgba(255, 255, 255, 0.06);
        padding: 25px;
        border-radius: 18px;
        box-shadow: 0 0 20px rgba(127, 92, 255, 0.2);
        border: 1px solid rgba(255, 255, 255, 0.1);
        backdrop-filter: blur(12px);
        margin-bottom: 25px;
    }

    .stButton>button {
        background: linear-gradient(90deg, #7f5cff, #4a2aff);
        color: white;
        border: none;
        padding: 12px 22px;
        border-radius: 10px;
        font-size: 16px;
        font-weight: 600;
        box-shadow: 0 0 15px #7f5cff;
        transition: 0.3s;
    }

    .stButton>button:hover {
        box-shadow: 0 0 25px #a78bff;
        transform: translateY(-2px);
    }

    .stTextInput>div>div>input, textarea {
        background: rgba(255, 255, 255, 0.08) !important;
        color: white !important;
        border-radius: 10px !important;
    }

    .stTextArea textarea {
        background: rgba(0, 0, 0, 0.3) !important;
        border: 1px solid rgba(127, 92, 255, 0.5);
        border-radius: 12px !important;
        color: #e3e3e3 !important;
    }

    </style>
""", unsafe_allow_html=True)

# ---------------- UI HEADER ----------------
st.markdown('<div class="title-glow">⚡ CO-Based Question Generator</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle"> Upload PDF → Choose CO → Generate Questions</div>', unsafe_allow_html=True)

# ---------------- INPUT UI ----------------
st.markdown('<div class="glass-card">', unsafe_allow_html=True)
subject_name = st.text_input("🎓 Subject Name")
uploaded_pdf = st.file_uploader("📄 Upload Syllabus PDF", type=["pdf"])
st.markdown('</div>', unsafe_allow_html=True)

if uploaded_pdf:
    st.session_state.syllabus_text = extract_text_from_pdf(uploaded_pdf)
    st.success("PDF extracted successfully!")

# ---------------- CO SELECTION ----------------
st.markdown('<div class="glass-card">', unsafe_allow_html=True)
st.subheader("🎯 Choose Course Outcome")

co_unit_map = {
    "CO1": 1,
    "CO2": 2,
    "CO3": 3,
    "CO4": 4,
    "CO5": 5
}

col1, col2, col3, col4, col5 = st.columns(5)

if col1.button("CO1"): st.session_state.selected_co, st.session_state.unit = "CO1", 1
if col2.button("CO2"): st.session_state.selected_co, st.session_state.unit = "CO2", 2
if col3.button("CO3"): st.session_state.selected_co, st.session_state.unit = "CO3", 3
if col4.button("CO4"): st.session_state.selected_co, st.session_state.unit = "CO4", 4
if col5.button("CO5"): st.session_state.selected_co, st.session_state.unit = "CO5", 5

st.markdown('</div>', unsafe_allow_html=True)

# ---------------- QUESTION COUNTS ----------------
st.markdown('<div class="glass-card">', unsafe_allow_html=True)
st.subheader("📝 Number of Questions")

colA, colB, colC = st.columns(3)
num_mcq = colA.slider("MCQs", 0, 50, 10)
num_short = colB.slider("Short Questions", 0, 50, 5)
num_long = colC.slider("Long Questions", 0, 20, 3)

st.markdown('</div>', unsafe_allow_html=True)

# ---------------- EXAMPLE FORMATS ----------------
st.markdown('<div class="glass-card">', unsafe_allow_html=True)
st.subheader("📌 Example Formats (Optional)")

example_mcq = st.text_area("Example MCQ")
example_short = st.text_area("Example Short Question")
example_long = st.text_area("Example Long Question")

st.markdown('</div>', unsafe_allow_html=True)

# ---------------- GENERATE BUTTON ----------------
if st.session_state.selected_co:
    st.info(f"Generating questions for {st.session_state.selected_co} (Unit {st.session_state.unit})")

if st.button("⚡ Generate Questions"):
    if not subject_name or not st.session_state.syllabus_text:
        st.error("❗ Subject name and syllabus PDF are required.")
    else:
        st.session_state.mcq = generate_questions(
            "multiple choice",
            subject_name,
            st.session_state.syllabus_text,
            st.session_state.unit,
            st.session_state.selected_co,
            num_mcq,
            example_mcq
        )

        st.session_state.short = generate_questions(
            "short answer",
            subject_name,
            st.session_state.syllabus_text,
            st.session_state.unit,
            st.session_state.selected_co,
            num_short,
            example_short
        )

        st.session_state.long = generate_questions(
            "long answer",
            subject_name,
            st.session_state.syllabus_text,
            st.session_state.unit,
            st.session_state.selected_co,
            num_long,
            example_long
        )

        st.success("🎉 Questions generated successfully!")

# ---------------- DISPLAY OUTPUT ----------------
st.markdown('<div class="glass-card">', unsafe_allow_html=True)
st.header("📄 Generated Questions")

if st.session_state.mcq:
    st.subheader("MCQs")
    st.text_area("MCQs", st.session_state.mcq, height=250)

if st.session_state.short:
    st.subheader("Short Questions")
    st.text_area("Short Questions", st.session_state.short, height=250)

if st.session_state.long:
    st.subheader("Long Questions")
    st.text_area("Long Questions", st.session_state.long, height=250)

st.markdown('</div>', unsafe_allow_html=True)

# ---------------- DOCX DOWNLOAD ----------------
if st.button("📥 Download DOCX"):
    if st.session_state.mcq or st.session_state.short or st.session_state.long:
        doc = docx.Document()

        doc.add_heading(subject_name, level=1)
        doc.add_paragraph(
            f"Course Outcome: {st.session_state.selected_co}\n"
            f"Unit: {st.session_state.unit}"
        )

        if st.session_state.mcq:
            doc.add_heading("MCQs", level=2)
            for line in st.session_state.mcq.split("\n"):
                doc.add_paragraph(line)

        if st.session_state.short:
            doc.add_heading("Short Answer Questions", level=2)
            for line in st.session_state.short.split("\n"):
                doc.add_paragraph(line)

        if st.session_state.long:
            doc.add_heading("Long Answer Questions", level=2)
            for line in st.session_state.long.split("\n"):
                doc.add_paragraph(line)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            "⬇️ Download Question Paper",
            buffer,
            file_name=f"{st.session_state.selected_co}_Question_Paper.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
