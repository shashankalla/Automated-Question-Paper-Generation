# export.py

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

def export_pdf(paper, filename="question_paper.pdf"):
    
    doc = SimpleDocTemplate(filename)
    elements = []
    styles = getSampleStyleSheet()
    
    elements.append(Paragraph(f"Subject: {paper['subject']}", styles["Heading1"]))
    elements.append(Spacer(1, 0.3 * inch))
    
    for section, questions in paper["sections"].items():
        elements.append(Paragraph(section, styles["Heading2"]))
        
        for i, q in enumerate(questions, 1):
            elements.append(Paragraph(f"{i}. {q['question']}", styles["Normal"]))
            elements.append(Spacer(1, 0.2 * inch))
    
    doc.build(elements)


def export_docx(paper, filename="question_paper.docx"):
    
    doc = Document()
    doc.add_heading(f"Subject: {paper['subject']}", level=1)
    
    for section, questions in paper["sections"].items():
        doc.add_heading(section, level=2)
        
        for i, q in enumerate(questions, 1):
            doc.add_paragraph(f"{i}. {q['question']}")
    
    doc.save(filename)